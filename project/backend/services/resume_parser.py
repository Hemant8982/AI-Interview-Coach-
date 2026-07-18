"""Resume parser service.

Extracts raw text from PDF and DOCX resume uploads.
"""

import io
import os

import PyPDF2
from docx import Document


class ResumeParser:
    """Parses uploaded resume files into plain text."""

    SUPPORTED_EXTENSIONS = {".pdf", ".docx"}

    @staticmethod
    def parse(file_stream, filename: str) -> str:
        """Parse a file stream into text based on extension.

        Args:
            file_stream: A binary file-like object.
            filename: Original filename used to detect format.

        Returns:
            Extracted plain text.
        """
        ext = os.path.splitext(filename)[1].lower()
        if ext == ".pdf":
            return ResumeParser._parse_pdf(file_stream)
        if ext == ".docx":
            return ResumeParser._parse_docx(file_stream)
        raise ValueError(f"Unsupported file type: {ext}. Please upload PDF or DOCX.")

    @staticmethod
    def _parse_pdf(file_stream) -> str:
        reader = PyPDF2.PdfReader(file_stream)
        text_parts = []
        for page in reader.pages:
            extracted = page.extract_text()
            if extracted:
                text_parts.append(extracted)
        return "\n".join(text_parts).strip()

    @staticmethod
    def _parse_docx(file_stream) -> str:
        doc = Document(io.BytesIO(file_stream.read()))
        text_parts = [para.text for para in doc.paragraphs if para.text.strip()]
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text.strip())
        return "\n".join(text_parts).strip()
